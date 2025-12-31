#!/usr/bin/env python3
"""
HTML to DOCX Converter with CSS Preservation
Converts ebook.html to DOCX format while preserving CSS styling
"""

import os
import re
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import cssutils
import logging

# Suppress cssutils warnings
cssutils.log.setLevel(logging.ERROR)

class HTMLToDOCXConverter:
    def __init__(self, html_file_path, output_path=None):
        self.html_file_path = Path(html_file_path)
        self.output_path = output_path or self.html_file_path.parent / "output" / "ebook.docx"
        self.doc = Document()
        self.styles = {}
        self.current_list_level = 0
        
    def parse_css_color(self, color_str):
        """Parse CSS color string to RGB values"""
        if not color_str:
            return None
            
        # Handle hex colors
        if color_str.startswith('#'):
            hex_color = color_str[1:]
            if len(hex_color) == 3:
                hex_color = ''.join([c*2 for c in hex_color])
            if len(hex_color) == 6:
                try:
                    return RGBColor(
                        int(hex_color[0:2], 16),
                        int(hex_color[2:4], 16),
                        int(hex_color[4:6], 16)
                    )
                except ValueError:
                    return None
        
        # Handle named colors
        color_map = {
            'black': RGBColor(0, 0, 0),
            'white': RGBColor(255, 255, 255),
            'red': RGBColor(255, 0, 0),
            'green': RGBColor(0, 128, 0),
            'blue': RGBColor(0, 0, 255),
            'navy': RGBColor(0, 0, 128),
            'teal': RGBColor(0, 128, 128),
            'gray': RGBColor(128, 128, 128),
            'silver': RGBColor(192, 192, 192),
        }
        
        # Custom colors from your CSS
        if color_str == '#2c3e50':
            return RGBColor(44, 62, 80)
        elif color_str == '#3498db':
            return RGBColor(52, 152, 219)
        elif color_str == '#2980b9':
            return RGBColor(41, 128, 185)
        elif color_str == '#7f8c8d':
            return RGBColor(127, 140, 141)
        
        return color_map.get(color_str.lower())
    
    def parse_css_size(self, size_str):
        """Parse CSS size string to points"""
        if not size_str:
            return None
            
        # Handle point sizes
        if size_str.endswith('pt'):
            try:
                return Pt(float(size_str[:-2]))
            except ValueError:
                return None
        
        # Handle pixel sizes (approximate conversion)
        elif size_str.endswith('px'):
            try:
                px_value = float(size_str[:-2])
                return Pt(px_value * 0.75)  # Approximate px to pt conversion
            except ValueError:
                return None
        
        # Handle em sizes (relative to base font size)
        elif size_str.endswith('em'):
            try:
                em_value = float(size_str[:-2])
                return Pt(em_value * 12)  # Assuming 12pt base font
            except ValueError:
                return None
        
        return None
    
    def extract_css_styles(self, html_content):
        """Extract and parse CSS styles from HTML"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Find all style tags
        style_tags = soup.find_all('style')
        css_content = '\n'.join([tag.get_text() for tag in style_tags])
        
        # Parse CSS
        sheet = cssutils.parseString(css_content)
        
        for rule in sheet:
            if rule.type == rule.STYLE_RULE:
                selector = rule.selectorText
                styles = {}
                
                for prop in rule.style:
                    styles[prop.name] = prop.value
                
                self.styles[selector] = styles
        
        return soup
    
    def create_custom_styles(self):
        """Create custom styles in the document"""
        # Create heading styles
        try:
            h1_style = self.doc.styles['Heading 1']
            h1_style.font.size = Pt(20)
            h1_style.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
            h1_style.font.bold = True
        except:
            pass
        
        try:
            h2_style = self.doc.styles['Heading 2']
            h2_style.font.size = Pt(16)
            h2_style.font.color.rgb = RGBColor(52, 152, 219)  # #3498db
            h2_style.font.bold = True
        except:
            pass
        
        try:
            h3_style = self.doc.styles['Heading 3']
            h3_style.font.size = Pt(14)
            h3_style.font.color.rgb = RGBColor(41, 128, 185)  # #2980b9
            h3_style.font.bold = True
        except:
            pass
        
        # Create code style
        try:
            code_style = self.doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(9)
        except:
            pass
        
        # Create code block style
        try:
            code_block_style = self.doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_block_style.font.name = 'Courier New'
            code_block_style.font.size = Pt(9)
            code_block_style.paragraph_format.left_indent = Inches(0.5)
            code_block_style.paragraph_format.space_before = Pt(6)
            code_block_style.paragraph_format.space_after = Pt(6)
        except:
            pass
    
    def process_element(self, element, parent_paragraph=None):
        """Process HTML element and convert to DOCX"""
        if element.name is None:
            # Text node
            text = element.string
            if text and text.strip():
                if parent_paragraph:
                    run = parent_paragraph.add_run(text)
                    return run
                else:
                    p = self.doc.add_paragraph(text)
                    return p
            return None
        
        # Handle different HTML elements
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            text = element.get_text().strip()
            
            if level == 1:
                p = self.doc.add_heading(text, level=1)
            elif level == 2:
                p = self.doc.add_heading(text, level=2)
            elif level == 3:
                p = self.doc.add_heading(text, level=3)
            else:
                p = self.doc.add_heading(text, level=4)
            
            return p
        
        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                p = self.doc.add_paragraph()
                self.process_inline_elements(element, p)
                return p
        
        elif element.name == 'pre':
            text = element.get_text()
            if text.strip():
                p = self.doc.add_paragraph(text)
                try:
                    p.style = 'Code Block'
                except:
                    # Fallback formatting
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                return p
        
        elif element.name == 'code':
            text = element.get_text()
            if parent_paragraph:
                run = parent_paragraph.add_run(text)
                try:
                    run.style = 'Code'
                except:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                return run
        
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = li.get_text().strip()
                if text:
                    p = self.doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
        
        elif element.name == 'table':
            self.process_table(element)
        
        elif element.name == 'strong' or element.name == 'b':
            text = element.get_text()
            if parent_paragraph:
                run = parent_paragraph.add_run(text)
                run.font.bold = True
                return run
        
        elif element.name == 'em' or element.name == 'i':
            text = element.get_text()
            if parent_paragraph:
                run = parent_paragraph.add_run(text)
                run.font.italic = True
                return run
        
        elif element.name == 'br':
            if parent_paragraph:
                parent_paragraph.add_run('\n')
        
        elif element.name == 'div':
            # Handle div elements by processing their children
            for child in element.children:
                self.process_element(child)
        
        # Process child elements
        if element.name not in ['pre', 'code', 'strong', 'b', 'em', 'i']:
            for child in element.children:
                self.process_element(child, parent_paragraph)
        
        return None
    
    def process_inline_elements(self, element, paragraph):
        """Process inline elements within a paragraph"""
        for child in element.children:
            if child.name is None:
                # Text node
                text = str(child).strip()
                if text:
                    paragraph.add_run(text)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.font.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.font.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                try:
                    run.style = 'Code'
                except:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            elif child.name == 'a':
                # Handle links
                text = child.get_text()
                href = child.get('href', '')
                if href:
                    run = paragraph.add_run(f"{text} ({href})")
                    run.font.color.rgb = RGBColor(52, 152, 219)  # Blue color
                else:
                    run = paragraph.add_run(text)
            elif child.name == 'br':
                paragraph.add_run('\n')
            else:
                # Recursively process other inline elements
                self.process_inline_elements(child, paragraph)
    
    def process_table(self, table_element):
        """Process HTML table and convert to DOCX table"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Count columns
        max_cols = 0
        for row in rows:
            cols = len(row.find_all(['td', 'th']))
            max_cols = max(max_cols, cols)
        
        if max_cols == 0:
            return
        
        # Create table
        docx_table = self.doc.add_table(rows=len(rows), cols=max_cols)
        docx_table.style = 'Table Grid'
        
        # Fill table
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    docx_cell = docx_table.cell(row_idx, col_idx)
                    docx_cell.text = cell.get_text().strip()
                    
                    # Bold header cells
                    if cell.name == 'th':
                        for paragraph in docx_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
    
    def convert(self):
        """Main conversion method"""
        print(f"🔄 Converting {self.html_file_path} to DOCX...")
        
        # Ensure output directory exists
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Read HTML file
        try:
            with open(self.html_file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        except Exception as e:
            print(f"❌ Error reading HTML file: {e}")
            return False
        
        # Parse HTML and extract CSS
        soup = self.extract_css_styles(html_content)
        
        # Create custom styles
        self.create_custom_styles()
        
        # Set document properties
        self.doc.core_properties.title = "Prompt Engineering for Developers"
        self.doc.core_properties.author = "Sunny Shivam"
        self.doc.core_properties.subject = "Crafting Intelligent LLM Solutions"
        
        # Process HTML body
        body = soup.find('body')
        if body:
            # Skip style tags and process content
            for element in body.children:
                if element.name != 'style':
                    self.process_element(element)
        
        # Save document
        try:
            self.doc.save(self.output_path)
            print(f"✅ Successfully converted to DOCX: {self.output_path}")
            return True
        except Exception as e:
            print(f"❌ Error saving DOCX file: {e}")
            return False

def install_dependencies():
    """Install required dependencies if not available"""
    required_packages = [
        'python-docx',
        'beautifulsoup4',
        'cssutils',
        'lxml'
    ]
    
    missing_packages = []
    
    for package in required_packages:
        try:
            if package == 'python-docx':
                import docx
            elif package == 'beautifulsoup4':
                import bs4
            elif package == 'cssutils':
                import cssutils
            elif package == 'lxml':
                import lxml
        except ImportError:
            missing_packages.append(package)
    
    if missing_packages:
        print(f"📦 Installing missing packages: {', '.join(missing_packages)}")
        import subprocess
        import sys
        
        for package in missing_packages:
            try:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
                print(f"✅ Installed {package}")
            except subprocess.CalledProcessError as e:
                print(f"❌ Failed to install {package}: {e}")
                return False
    
    return True

def main():
    """Main function"""
    print("🔄 HTML to DOCX Converter")
    print("=" * 40)
    
    # Check and install dependencies
    if not install_dependencies():
        print("❌ Failed to install required dependencies")
        return
    
    # Define paths
    script_dir = Path(__file__).parent
    html_file = script_dir / "htmlFiles" / "ebook.html"
    output_file = script_dir / "output" / "ebook.docx"
    
    # Check if HTML file exists
    if not html_file.exists():
        print(f"❌ HTML file not found: {html_file}")
        return
    
    # Convert
    converter = HTMLToDOCXConverter(html_file, output_file)
    success = converter.convert()
    
    if success:
        file_size = output_file.stat().st_size / (1024 * 1024)
        print(f"📄 File size: {file_size:.2f} MB")
        print(f"🎉 Conversion complete!")
        print(f"📖 Open the DOCX file with Microsoft Word or LibreOffice Writer")
    else:
        print("❌ Conversion failed")

if __name__ == "__main__":
    main()
